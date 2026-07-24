"""Markdown report body -> Word .docx.

The report content is a known, narrow Markdown subset — two `##` sections,
paragraphs, bullet and numbered lists, and `**bold**` — the same shape the
frontend renderer handles. So this is a small, deliberate parser for that
subset, not a general Markdown engine: a full converter (pandoc, a markdown
library plus a docx writer) is real weight for output we control end to end.

Anything unrecognised degrades to a plain paragraph rather than being dropped,
so an unexpected construct still appears in the document.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt
from docx.text.paragraph import Paragraph

_BULLET = re.compile(r"^[-*]\s+(.*)$")
_NUMBERED = re.compile(r"^\d+\.\s+(.*)$")
_HEADING = re.compile(r"^(#{1,4})\s+(.*)$")
# **bold** spans. Non-nesting, which the report prose never needs.
_BOLD = re.compile(r"\*\*([^*]+)\*\*")


def _add_runs(paragraph: Paragraph, text: str) -> None:
    """Split on `**bold**` and add each span as a run, bold where marked."""
    cursor = 0
    for match in _BOLD.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        paragraph.add_run(match.group(1)).bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _render_body(document: DocxDocument, content: str) -> None:
    # A paragraph is one or more consecutive non-blank, non-list lines; a blank
    # line or a structural line flushes it.
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            _add_runs(document.add_paragraph(), " ".join(paragraph))
            paragraph.clear()

    for raw in content.replace("\r\n", "\n").split("\n"):
        line = raw.rstrip()

        heading = _HEADING.match(line)
        bullet = _BULLET.match(line)
        numbered = _NUMBERED.match(line)

        if heading:
            flush()
            # `##` -> Heading 2, matching the report's two top-level sections;
            # deeper markers step down but never above Heading 2.
            level = max(2, min(len(heading.group(1)) + 1, 4))
            document.add_heading(heading.group(2), level=level)
        elif bullet:
            flush()
            _add_runs(document.add_paragraph(style="List Bullet"), bullet.group(1))
        elif numbered:
            flush()
            _add_runs(document.add_paragraph(style="List Number"), numbered.group(1))
        elif line.strip() == "":
            flush()
        else:
            paragraph.append(line)

    flush()


def report_to_docx(
    *, title: str, strategy: str | None, generated_at: str, content: str
) -> bytes:
    """Render a report to a .docx and return the file bytes.

    The provenance line (strategy + timestamp) is kept in the document, not just
    the filename: a report exported for a dissertation appendix has to carry
    which prompting strategy produced it, so a grounded and a naive export are
    never confused.
    """
    document = Document()

    normal = document.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(11)

    document.add_heading(title, level=0)

    meta = document.add_paragraph()
    provenance = f"Prompting strategy: {strategy or 'n/a'}  ·  Generated: {generated_at}"
    meta.add_run(provenance).italic = True

    _render_body(document, content)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
